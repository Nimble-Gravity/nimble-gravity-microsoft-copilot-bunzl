import os
import subprocess
import sys
import unittest

import openpyxl
from docx import Document
from pptx import Presentation

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.normpath(os.path.join(HERE, '..', '..', 'assets', 'lab-data'))

XLSX_PATH = os.path.join(OUT_DIR, 'bunzl-quarterly-budget-review.xlsx')
PPTX_PATH = os.path.join(OUT_DIR, 'bunzl-business-review.pptx')
DOCX_PATH = os.path.join(OUT_DIR, 'bunzl-team-update-memo.docx')


class TestGeneratedFiles(unittest.TestCase):
    def test_xlsx_has_expected_sheets_and_row_count(self):
        wb = openpyxl.load_workbook(XLSX_PATH)
        self.assertEqual(wb.sheetnames, ['Budget vs Actual', 'Notes'])
        sheet = wb['Budget vs Actual']
        # header row + 5 segment rows
        self.assertEqual(sheet.max_row, 6)
        self.assertEqual(sheet['A1'].value, 'Segment')

    def test_pptx_has_expected_slide_count(self):
        prs = Presentation(PPTX_PATH)
        self.assertEqual(len(prs.slides), 4)

    def test_docx_has_expected_title_paragraph(self):
        doc = Document(DOCX_PATH)
        self.assertGreater(len(doc.paragraphs), 0)
        self.assertIn('Team Update', doc.paragraphs[0].text)


if __name__ == '__main__':
    unittest.main()
